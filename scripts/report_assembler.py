#!/usr/bin/env python3
"""
报告拼装器 - 纯拼装，零 LLM 调用

设计原则：
- 文本内容全部由 Claude（调用方）提供
- 本模块只负责：结构化落盘、docx 拼装、文件齐全度自检
- 严格拒绝任何"自动生成占位文本""模板套话"

接口分工：
- save_insight: Claude 写完一段业务洞察后落盘（强制 4 要素：现象/模式/根因/影响）
- save_recommendation: Claude 写完一条建议后落盘（强制 5 要素：动作/对象/方法/收益/周期）
- assemble_report: 把 Claude 写的所有 sections 拼成 docx
"""

import os
import json
import re
from datetime import datetime
from typing import Dict, Any, List, Optional


# 4 要素强约束（防止 Claude 偷懒写空话）
INSIGHT_REQUIRED = ["phenomenon", "pattern", "root_cause", "impact"]
INSIGHT_LABELS = {
    "phenomenon": "现象",
    "pattern": "模式",
    "root_cause": "根因",
    "impact": "影响",
}

# 5 要素强约束
RECOMMENDATION_REQUIRED = ["action", "target", "method", "gain", "cycle"]
RECOMMENDATION_LABELS = {
    "action": "动作",
    "target": "对象",
    "method": "方法",
    "gain": "收益",
    "cycle": "周期",
}

# 空话黑名单（命中即拒绝）
BANNED_PHRASES = [
    "数据质量良好", "建议使用模型监测", "需要进一步分析",
    "准确率较高", "效果较好", "整体表现不错", "结果令人满意",
    "需要更多数据", "请进一步研究", "建议持续优化",
]


def save_insight(base_dir: str, step: str, insight: Dict[str, str],
                 strict: bool = True) -> Dict[str, Any]:
    """
    保存一条业务洞察。

    参数:
        base_dir: outputs/{model}/{dataset}/ 的根目录
        step: 步骤名（如 'step2_eda', 'step4_anomaly'）
        insight: {
            "title": "洞察标题",
            "phenomenon": "观察到的现象（必须包含具体数值）",
            "pattern": "数据模式/规律",
            "root_cause": "可能的根因",
            "impact": "对业务的影响",
            "evidence": ["证据1", "证据2"],  # 可选，引用的文件路径
        }
        strict: 是否启用 4 要素强约束（默认 True）

    返回:
        {"status": "ok"/"rejected", "path": ..., "issues": [...]}
    """
    issues = _validate_insight(insight, strict)
    if issues and strict:
        return {
            "status": "rejected",
            "issues": issues,
            "hint": "Claude 请重写：4要素必须实质性填充，不能空话",
        }

    step_dir = os.path.join(base_dir, step)
    os.makedirs(step_dir, exist_ok=True)

    insights_path = os.path.join(step_dir, "insights.json")
    existing = []
    if os.path.exists(insights_path):
        with open(insights_path, "r", encoding="utf-8") as f:
            existing = json.load(f)

    insight_record = {
        "id": f"INS-{len(existing)+1:03d}",
        "timestamp": datetime.now().isoformat(),
        "title": insight.get("title", ""),
        "phenomenon": insight.get("phenomenon", ""),
        "pattern": insight.get("pattern", ""),
        "root_cause": insight.get("root_cause", ""),
        "impact": insight.get("impact", ""),
        "evidence": insight.get("evidence", []),
    }
    existing.append(insight_record)

    with open(insights_path, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)

    return {
        "status": "ok",
        "path": insights_path,
        "id": insight_record["id"],
        "warnings": issues if not strict else [],
    }


def save_recommendation(base_dir: str, step: str, rec: Dict[str, str],
                        strict: bool = True) -> Dict[str, Any]:
    """
    保存一条业务建议（5 要素：动作/对象/方法/收益/周期）。

    参数:
        rec: {
            "title": "建议标题",
            "action": "应执行的动作（动词开头）",
            "target": "针对哪个设备/工艺/参数",
            "method": "具体怎么做（含参数）",
            "gain": "预期收益（数值化）",
            "cycle": "实施周期（如 2周/1月）",
        }
    """
    issues = _validate_recommendation(rec, strict)
    if issues and strict:
        return {
            "status": "rejected",
            "issues": issues,
            "hint": "Claude 请重写：5要素必须有实质内容，收益必须数值化",
        }

    step_dir = os.path.join(base_dir, step)
    os.makedirs(step_dir, exist_ok=True)

    recs_path = os.path.join(step_dir, "recommendations.json")
    existing = []
    if os.path.exists(recs_path):
        with open(recs_path, "r", encoding="utf-8") as f:
            existing = json.load(f)

    rec_record = {
        "id": f"REC-{len(existing)+1:03d}",
        "timestamp": datetime.now().isoformat(),
        "title": rec.get("title", ""),
        "action": rec.get("action", ""),
        "target": rec.get("target", ""),
        "method": rec.get("method", ""),
        "gain": rec.get("gain", ""),
        "cycle": rec.get("cycle", ""),
    }
    existing.append(rec_record)

    with open(recs_path, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)

    return {"status": "ok", "path": recs_path, "id": rec_record["id"]}


def assemble_report(base_dir: str, sections: List[Dict[str, Any]],
                    title: str = "工业数据分析报告",
                    output_format: str = "docx") -> Dict[str, Any]:
    """
    把 Claude 写好的 sections 拼装成 docx 报告。

    本工具不做任何"内容生成"——所有文本由 Claude 提供。

    参数:
        sections: [
            {
                "title": "章节标题",
                "content_md": "Markdown 格式的章节正文（Claude 写）",
                "figures": ["相对于 base_dir 的图片路径", ...],  # 可选
                "tables": [{"caption": "表名", "data": [[...]]}],  # 可选
            },
            ...
        ]
        output_format: 'docx' | 'md'

    返回:
        {"status": "ok", "path": "..."}
    """
    if not sections:
        return {"status": "error", "msg": "sections 为空，Claude 必须提供内容"}

    # 检查每个 section 是否有实质内容
    rejected = []
    for i, sec in enumerate(sections):
        content = sec.get("content_md", "").strip()
        if not content or len(content) < 30:
            rejected.append(f"section[{i}] '{sec.get('title','')}'  内容过短或为空")
        for banned in BANNED_PHRASES:
            if banned in content:
                rejected.append(f"section[{i}] 命中空话黑名单: '{banned}'")
    if rejected:
        return {
            "status": "rejected",
            "issues": rejected,
            "hint": "Claude 请重写章节内容，禁止使用空话黑名单",
        }

    os.makedirs(base_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if output_format == "md":
        return _assemble_markdown(base_dir, sections, title, timestamp)
    return _assemble_docx(base_dir, sections, title, timestamp)


def _assemble_markdown(base_dir: str, sections: List[Dict],
                        title: str, timestamp: str) -> Dict[str, Any]:
    out_path = os.path.join(base_dir, f"report_{timestamp}.md")
    lines = [f"# {title}\n", f"_生成时间: {datetime.now().isoformat()}_\n"]
    for sec in sections:
        lines.append(f"\n## {sec.get('title','')}\n")
        lines.append(sec.get("content_md", "") + "\n")
        for fig in sec.get("figures", []):
            fig_rel = os.path.relpath(fig, base_dir) if os.path.isabs(fig) else fig
            lines.append(f"\n![]({fig_rel})\n")
        for tbl in sec.get("tables", []):
            lines.append(_table_to_md(tbl))
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return {"status": "ok", "path": out_path, "format": "md",
            "n_sections": len(sections)}


def _assemble_docx(base_dir: str, sections: List[Dict],
                    title: str, timestamp: str) -> Dict[str, Any]:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        # 降级到 markdown
        return _assemble_markdown(base_dir, sections, title, timestamp)

    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    n_figures = 0
    for sec in sections:
        doc.add_heading(sec.get("title", ""), level=1)

        # 渲染 markdown 段落（简单处理）
        for para in sec.get("content_md", "").split("\n"):
            para = para.strip()
            if not para:
                continue
            if para.startswith("### "):
                doc.add_heading(para[4:], level=3)
            elif para.startswith("## "):
                doc.add_heading(para[3:], level=2)
            elif para.startswith("- "):
                doc.add_paragraph(para[2:], style="List Bullet")
            else:
                doc.add_paragraph(para)

        for fig in sec.get("figures", []):
            fig_path = fig if os.path.isabs(fig) else os.path.join(base_dir, fig)
            if os.path.exists(fig_path):
                try:
                    doc.add_picture(fig_path, width=Inches(5.5))
                    n_figures += 1
                except Exception:
                    pass

        for tbl in sec.get("tables", []):
            _add_docx_table(doc, tbl)

    out_path = os.path.join(base_dir, f"report_{timestamp}.docx")
    doc.save(out_path)
    return {
        "status": "ok",
        "path": out_path,
        "format": "docx",
        "n_sections": len(sections),
        "n_figures_embedded": n_figures,
    }


def _add_docx_table(doc, tbl: Dict):
    data = tbl.get("data", [])
    if not data:
        return
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i, j).text = str(cell)
    if tbl.get("caption"):
        doc.add_paragraph(tbl["caption"], style="Caption")


def _table_to_md(tbl: Dict) -> str:
    data = tbl.get("data", [])
    if not data:
        return ""
    lines = []
    if tbl.get("caption"):
        lines.append(f"\n**表 {tbl['caption']}**\n")
    header = data[0]
    lines.append("| " + " | ".join(str(c) for c in header) + " |")
    lines.append("|" + "|".join(["---"] * len(header)) + "|")
    for row in data[1:]:
        lines.append("| " + " | ".join(str(c) for c in row) + " |")
    return "\n".join(lines) + "\n"


# ==================== 校验函数 ====================

def _validate_insight(insight: Dict, strict: bool) -> List[str]:
    issues = []
    for key in INSIGHT_REQUIRED:
        val = insight.get(key, "").strip()
        if not val:
            issues.append(f"缺失要素 [{INSIGHT_LABELS[key]}={key}]")
        elif len(val) < 8:
            issues.append(f"要素 [{INSIGHT_LABELS[key]}] 内容过短(<8字)")
        for banned in BANNED_PHRASES:
            if banned in val:
                issues.append(f"[{INSIGHT_LABELS[key]}] 命中空话: '{banned}'")
    # phenomenon 必须含数字（事实必须数值化）
    phenom = insight.get("phenomenon", "")
    if phenom and not re.search(r"\d", phenom):
        issues.append("[现象] 必须包含具体数值（如百分比、计数、阈值）")
    return issues


def _validate_recommendation(rec: Dict, strict: bool) -> List[str]:
    issues = []
    for key in RECOMMENDATION_REQUIRED:
        val = rec.get(key, "").strip()
        if not val:
            issues.append(f"缺失要素 [{RECOMMENDATION_LABELS[key]}={key}]")
    # gain 必须含数字（收益必须数值化）
    gain = rec.get("gain", "")
    if gain and not re.search(r"\d", gain):
        issues.append("[收益] 必须包含具体数值（如降本%、提效%）")
    # cycle 必须含时间单位
    cycle = rec.get("cycle", "")
    if cycle and not re.search(r"(天|周|月|年|day|week|month|year|h|min)", cycle):
        issues.append("[周期] 必须包含时间单位")
    return issues


# ==================== 自检 ====================

def quality_check(base_dir: str) -> Dict[str, Any]:
    """
    自检报告齐全度。

    检查项：
    - 各 step 目录是否存在
    - insights.json / recommendations.json 是否齐全
    - 4/5 要素是否填充
    - 是否有报告产物（docx/md）
    """
    result = {"base_dir": base_dir, "issues": [], "summary": {}}

    if not os.path.isdir(base_dir):
        result["issues"].append(f"目录不存在: {base_dir}")
        return result

    # 收集所有 insight/rec
    n_insights = 0
    n_recs = 0
    weak_insights = []
    for step_dir in os.listdir(base_dir):
        step_path = os.path.join(base_dir, step_dir)
        if not os.path.isdir(step_path):
            continue

        ins_path = os.path.join(step_path, "insights.json")
        if os.path.exists(ins_path):
            with open(ins_path, "r", encoding="utf-8") as f:
                ins = json.load(f)
                n_insights += len(ins)
                for i in ins:
                    issues = _validate_insight(i, strict=True)
                    if issues:
                        weak_insights.append({"id": i.get("id"), "issues": issues})

        rec_path = os.path.join(step_path, "recommendations.json")
        if os.path.exists(rec_path):
            with open(rec_path, "r", encoding="utf-8") as f:
                n_recs += len(json.load(f))

    has_report = any(f.endswith((".docx", ".md")) and "report" in f.lower()
                     for f in os.listdir(base_dir))

    result["summary"] = {
        "total_insights": n_insights,
        "total_recommendations": n_recs,
        "has_report": has_report,
        "weak_insights_count": len(weak_insights),
    }

    if n_insights < 5:
        result["issues"].append(f"洞察数量 {n_insights} < 5（红线要求 ≥ 5）")
    if n_recs < 3:
        result["issues"].append(f"建议数量 {n_recs} < 3（红线要求 ≥ 3）")
    if weak_insights:
        result["issues"].append(f"{len(weak_insights)} 条洞察未通过 4 要素校验")
        result["weak_insights_detail"] = weak_insights[:10]
    if not has_report:
        result["issues"].append("未发现报告文件（docx/md）")

    result["passed"] = len(result["issues"]) == 0
    return result
